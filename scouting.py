import pandas as pd
import gspread
from google.oauth2.service_account import Credentials
from shiny import App, render, ui, reactive
from shiny.ui import HTML
from datetime import datetime 
import pandas as pd
from shiny import App, reactive, render, ui
from htmltools import HTML
import json
from datetime import datetime
from pytube import Search
import requests
from bs4 import BeautifulSoup
import nest_asyncio
from shiny import run_app
from docx import Document
import matplotlib.pyplot as plt
import seaborn as sns
import os 

nest_asyncio.apply()

# Configuración de Google Sheets
SCOPES = ['https://www.googleapis.com/auth/spreadsheets']

# Cargar la variable de entorno que contiene las credenciales de Google
creds_json = os.getenv("GOOGLE_APPLICATION_CREDENTIALS_JSON")

# Verifica si la variable fue cargada correctamente
if creds_json is None:
    print("Error: La variable de entorno no está configurada correctamente.")
else:
    print("La variable de entorno se ha cargado correctamente.")

# Convierte el JSON cargado en un diccionario y crea las credenciales
creds_dict = json.loads(creds_json)
# Crear las credenciales con el alcance adecuado
creds = Credentials.from_service_account_info(creds_dict, scopes=SCOPES)

# Autorizar el cliente de gspread con las credenciales
client = gspread.authorize(creds)

# Autoriza el cliente de gspread con las credenciales
client = gspread.authorize(creds)
SHEET_ID = "1MXuIF81o1Ts_QbEdhm0X790p4qRYxxtgoi3ufRwzEK8"

workbook = client.open_by_key(SHEET_ID)
worksheet = workbook.sheet1
values = worksheet.get_all_values()
df_jugadores = pd.DataFrame(values[1:], columns=values[0])
INFORMES_SHEET_NAME = "Historial"
LINK_SHEET_NAME = "Links"
def cargar_links():
    workbook = client.open_by_key(SHEET_ID)
    try:
        worksheet = workbook.worksheet(LINK_SHEET_NAME)
    except gspread.exceptions.WorksheetNotFound:
        return {}
    
    values = worksheet.get_all_values()
    if len(values) < 2:
        return {}
    
    df_links = pd.DataFrame(values[1:], columns=values[0])
    historial_links = {}
    
    for _, row in df_links.iterrows():
        jugador = row["Jugador"]
        link = row["Link"]
        
        if jugador not in historial_links:
            historial_links[jugador] = []
        historial_links[jugador].append(link)
    
    return historial_links

def guardar_links(historial_links):
    workbook = client.open_by_key(SHEET_ID)
    
    try:
        worksheet = workbook.worksheet(LINK_SHEET_NAME)
    except gspread.exceptions.WorksheetNotFound:
        worksheet = workbook.add_worksheet(title=LINK_SHEET_NAME, rows="1000", cols="2")
    
    df_links = pd.DataFrame([
        {"Jugador": jugador, "Link": link}
        for jugador, links in historial_links.items()
        for link in links
    ])
    
    values_actualizados = [df_links.columns.tolist()] + df_links.values.tolist()
    
    worksheet.clear()
    worksheet.update(values_actualizados)

    print("Historial de enlaces actualizado en Google Sheets.")
    
# 🔹 Función para cargar informes desde Google Sheets
def cargar_historial():
    workbook = client.open_by_key(SHEET_ID)
    worksheet = workbook.worksheet(INFORMES_SHEET_NAME)
    
    values = worksheet.get_all_values()
    if len(values) < 2:
        return {}

    df_informes = pd.DataFrame(values[1:], columns=values[0])
    historial = {}
    
    for _, row in df_informes.iterrows():
        jugador = row["Jugador"]
        informe = {"Fecha": row["Fecha"], "Título": row["Título"], "Texto": row["Texto"]}
        
        if jugador not in historial:
            historial[jugador] = []
        historial[jugador].append(informe)

    return historial

# 🔹 Función para guardar informes en Google Sheets
def guardar_historial(historial):
    workbook = client.open_by_key(SHEET_ID)
    
    try:
        worksheet = workbook.worksheet(INFORMES_SHEET_NAME)
    except gspread.exceptions.WorksheetNotFound:
        worksheet = workbook.add_worksheet(title=INFORMES_SHEET_NAME, rows="1000", cols="4")

    df_historial = pd.DataFrame([
        {"Jugador": jugador, **informe}
        for jugador, informes in historial.items()
        for informe in informes
    ])

    values_actualizados = [df_historial.columns.tolist()] + df_historial.values.tolist()

    worksheet.clear()
    worksheet.update(values_actualizados)

    print("Historial de informes actualizado en Google Sheets.")

# 🔹 Estado reactivo para historial
historial_data = reactive.Value(cargar_historial())

# 🔹 Función para guardar un informe
def guardar_informe():
    jugador = input.jugador_informe()
    titulo = input.titulo_informe()
    texto = input.texto_informe()

    if not jugador or not titulo or not texto:
        return  

    fecha = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    nuevo_informe = {"Fecha": fecha, "Título": titulo, "Texto": texto}

    historial_actual = historial_data.get().copy()
    historial_actual.setdefault(jugador, []).append(nuevo_informe)

    guardar_historial(historial_actual)
    historial_data.set(historial_actual)

    ui.update_text("titulo_informe", value="")
    ui.update_text("texto_informe", value="")
    
# Calcular la edad
current_year = datetime.now().year
import hashlib

def cargar_datos():
    """Función para obtener los datos más recientes de Google Sheets"""
    values = worksheet.get_all_values()
    df = pd.DataFrame(values[1:], columns=values[0])
    
    # Generar un hash del DataFrame para detectar cambios
    df_hash = hashlib.md5(pd.util.hash_pandas_object(df, index=True).values).hexdigest()
    
    return df, df_hash  # Retornamos el DataFrame y su hash

# Configurar el objeto reactivo con actualización automática
df_reactivo = reactive.Value(cargar_datos()[0])  # Inicializar con los datos actuales

# Función para limpiar y convertir las fechas
def clean_birthday(value):
    try:
        value = str(value).strip()  # Eliminar espacios en blanco
        if value.isdigit() and len(value) == 4:  # Si es un año (YYYY)
            return f"{value}-01-01"  # Convertirlo a 'YYYY-01-01'
        return pd.to_datetime(value, errors='coerce')  # Intentar convertir a fecha
    except:
        return pd.NaT


# Aplicamos la función a la columna 'Birthdate'
df_jugadores["Birthdate"] = df_jugadores["Birthdate"].apply(clean_birthday)
df_jugadores["Birthdate"] = pd.to_datetime(df_jugadores["Birthdate"], errors="coerce")
df_jugadores["year_of_birth"] = df_jugadores["Birthdate"].dt.year
df_jugadores["Birthdate"] = df_jugadores["Birthdate"].dt.strftime('%Y-%m-%d')
df_jugadores['Club Contract'] = pd.to_datetime(df_jugadores['Club Contract'], errors='coerce').dt.date
df_jugadores["Contacto"] = df_jugadores["Contacto"].replace("", "No")

contacto_colors = {
    "Si": "background-color: #49bb25 ; color: white;",  # Verde para "Sí"
    "No": "background-color: #dc3545; color: white;"   # Rojo para "No"
}

# Convertir la columna 'Birthdate' a datetime
df_jugadores['Birthdate'] = pd.to_datetime(df_jugadores['Birthdate'], errors='coerce').dt.date

# Calcular la edad de los jugadores
def calcular_edad(fecha_nacimiento):
    if pd.isna(fecha_nacimiento):
        return None
    hoy = datetime.today().date()
    edad = hoy.year - fecha_nacimiento.year - ((hoy.month, hoy.day) < (fecha_nacimiento.month, fecha_nacimiento.day))
    
    return edad

df_jugadores['Edad'] = df_jugadores['Birthdate'].apply(calcular_edad)
df_jugadores['Edad'] = df_jugadores['Edad'].fillna("").astype(str) 
df_jugadores['Birthdate'] = df_jugadores['Birthdate'].fillna("").astype(str) 
df_jugadores['Club Contract'] = df_jugadores['Club Contract'].fillna("").astype(str) 
df_jugadores['Edad'] = df_jugadores['Edad'].apply(lambda x: int(float(x)) if x != "" else "")

# Obtener lista de nacionalidades únicas, dividiendo combinaciones y ordenando
nacionalidades = sorted(set(sum([str(n).split(',') for n in df_jugadores["Nationality"].dropna().unique()], [])))

# Definir colores para cada valor en "Assessment"
assessment_colors = {
    "Seguir": "background-color: #BB252E; color: white;",
    "Llamar": "background-color: #17a2b8; color: white;",
    "NaN": "background-color: #ffc107; color: black;",
    "All Iron Sports": "background-color: #49BB25; color: white;"
}

# Convertir diccionario de colores en JSON para JavaScript
colors_json = json.dumps(assessment_colors)

# Filtrar y ordenar las opciones
jugadores_unicos = sorted(df_jugadores["full_name"].dropna().unique().tolist())
columnas_unicas = sorted(df_jugadores.columns.tolist())
posicion_1_unica = sorted(df_jugadores["position_1"].dropna().unique().tolist())
posicion_2_unica = sorted(df_jugadores["position_2"].dropna().unique().tolist())
equipos_unicos = sorted(df_jugadores["Team"].dropna().unique().tolist())
nacionalidades_unicas = sorted(set(sum([str(n).split(',') for n in df_jugadores["Nationality"].dropna().unique()], [])))
agencias_unicas = sorted(df_jugadores["Agency"].dropna().unique().tolist())
seguimiento_unico = sorted(df_jugadores["Assessment"].dropna().unique().tolist())
categoria_unica = sorted(df_jugadores["Categoria"].dropna().unique().tolist())

# Función para buscar equipo y jugador, devolviendo solo el primer enlace
def buscar_equipo_jugador(jugador, equipo):
    query = f"{jugador} {equipo} lapreferente"
    url = f"https://duckduckgo.com/html/?q={query}"
    headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"}
    response = requests.get(url, headers=headers)
    if response.status_code == 200:
        soup = BeautifulSoup(response.text, "html.parser")
        enlace = soup.find("a", class_="result__a", href=True)
        return enlace["href"] if enlace else None
    else:
        print("Error al realizar la solicitud")
        return None

def es_doble_nacionalidad(nacionalidad):
    return "," in nacionalidad  # Si hay una coma, es doble nacionalidad

def generar_tarjeta(titulo, datos):
    contenido = f'''
    <div class="card" style="margin-bottom: 20px; padding: 15px; border: 1px solid #ddd; border-radius: 8px; 
                             background-color: #f9f9f9; box-shadow: 2px 2px 10px rgba(0,0,0,0.1);">
        <h4 style="margin-top: 0; color: #333;">{titulo}</h4>
    '''
    if isinstance(datos, dict):
        for clave, valor in datos.items():
            contenido += f'<p><strong>{clave}:</strong> {valor}</p>'
    else:
        contenido += f'<p>{datos}</p>'
    
    contenido += '</div>'
    return contenido

columnas = [col for col in df_jugadores.columns if col != 'Comentarios'] + ['Comentarios']
df_jugadores = df_jugadores[columnas]
columnas_ordenables = [col for col in columnas_unicas if col != "Edad"]  # Excluir "Edad"

app_ui = ui.page_navbar(
    ui.nav_panel(
        "Tabla Jugadores",
        ui.layout_sidebar(
            ui.sidebar(
                ui.panel_well(
                    ui.input_selectize("filtro_categoria", "Seleccionar categoria", choices=categoria_unica, remove_button=True, selected="Juveniles"),
                    ui.input_text("filtro_nombre", "Filtrar por Nombre"),
                    ui.input_selectize("filtro_seguimiento", "Filtrar por Seguimiento", choices=[""] + seguimiento_unico, remove_button=True),
                    ui.input_selectize("filtro_posicion1", "Filtrar por Posición 1", choices=[""] + posicion_1_unica, remove_button=True),
                    ui.input_selectize("filtro_posicion2", "Filtrar por Posición 2", choices=[""] + posicion_2_unica, remove_button=True),
                    ui.input_selectize("filtro_equipo", "Filtrar por Equipo", choices=[""] + equipos_unicos, remove_button=True),
                    ui.input_selectize("filtro_nacionalidad", "Filtrar por Nacionalidad", choices=[""] + nacionalidades_unicas, remove_button=True),
                    ui.input_checkbox("doble_nacionalidad", "Mostrar solo jugadores con doble nacionalidad", False),
                    ui.input_selectize("filtro_agencia", "Filtrar por Agencia", choices=[""] + agencias_unicas, remove_button=True),
                    ui.input_slider("filtro_edad_min", "Edad Mínima", min=1998, max=2012, value=1998),
                    ui.input_slider("filtro_edad_max", "Edad Máxima", min=1999, max=2012, value=2012),
                    ui.input_text("filtro_comentarios", "Filtrar por Comentarios"),
                    ui.input_action_button("restablecer_filtros", "Restablecer Filtros"),
                )
            ),
            ui.panel_well(
                ui.layout_columns(
                    ui.input_selectize("jugador", "Seleccionar Jugador", choices=jugadores_unicos, remove_button=True),
                    ui.input_selectize("columna", "Seleccionar Columna", choices=columnas_unicas, remove_button=True, selected="Comentarios"),
                    ui.input_text("nuevo_valor", "Nuevo Valor"),
                    ui.input_action_button("actualizar", "Actualizar"),
                    ui.input_selectize("orden_columna", "Ordenar por columna", choices=columnas_ordenables, selected="full_name"),
                    ui.input_selectize("orden_direccion", "Orden", choices=["asc", "desc"], selected="asc"),
                    ui.input_action_button("ordenar", "Ordenar"),
                    col_widths=[3, 3, 3, 2],  # Ajusta el ancho de las columnas

                )
            ),
            ui.panel_well(
                ui.output_ui("tabla_actualizada")
            ),
            ui.panel_well(
                ui.layout_columns(
                    ui.h4("Agregar Nuevo Jugador"),
                    ui.input_text("nuevo_nombre", "Nombre"),
                    ui.input_selectize("nueva_posicion1", "Posición 1", choices=posicion_1_unica, remove_button=True),
                    ui.input_selectize("nueva_posicion2", "Posición 2", choices=posicion_2_unica, remove_button=True),
                    ui.input_date("nuevo_birthdate", "Fecha de Nacimiento"),
                    ui.input_selectize("nuevo_pie", "Pie", choices=["Derecho", "Izquierdo", "Ambidiestro"], remove_button=True),
                    ui.input_selectize("nuevo_equipo", "Equipo", choices=equipos_unicos, remove_button=True),
                    ui.input_selectize("nuevo_assessment", "Seguimiento", choices=seguimiento_unico, remove_button=True),
                    ui.input_selectize("nueva_nacionalidad", "Nacionalidad", choices=nacionalidades_unicas, remove_button=True),
                    ui.input_selectize("nueva_agencia", "Agencia", choices=agencias_unicas, remove_button=True),
                    ui.input_text("nuevo_club_contract", "Contrato con Club"),
                    ui.input_text("nuevo_contacto", "Contacto"),
                    ui.input_selectize("nueva_categoria", "Categoría", choices=categoria_unica, remove_button=True),
                    ui.input_text("nuevo_comentario", "Comentario"),
                    ui.input_action_button("agregar_jugador", "Agregar Jugador"))
)
        )
    ),
    ui.nav_panel(
        "Detalle Jugador",
        ui.layout_sidebar(
            ui.sidebar(
                ui.input_selectize("jugador_detalle", "Seleccionar Jugador", choices=jugadores_unicos, remove_button=True),
                ui.input_select("informe_seleccionado", "Seleccionar Informe para Editar", choices=["Nuevo Informe"], multiple=False,selected = "Nuevo Informe"),
                ui.input_text("titulo_informe", "Título del Informe"),
                ui.input_text_area("texto_informe", "Texto del Informe", placeholder="Escribe tu informe aquí..."),
                ui.input_action_button("guardar_informe", "Guardar Informe"),
                ui.input_action_button("eliminar_informe", "Eliminar Informe", style="background-color: red; color: white;"),
                ui.download_button("descargar_historial", "Descargar Historial"),
                ui.input_text("nuevo_enlace", "Añadir enlace"),  
                ui.input_action_button("guardar_enlace", "Guardar Enlace")
            ),
            ui.panel_well(
                ui.output_ui("detalle_jugador"),
                ui.output_ui("historial_informes"),
                ui.output_ui("lapreferente_jugador"),
                ui.output_ui("historial_enlaces")
                         
            )
        )
),
    ui.nav_panel(
        "Gráficos",
        ui.layout_sidebar(
            ui.sidebar(
                ui.input_slider("top_n", "Número de Nacionalidades Destacadas", min=2, max=10, value=3),
            ),
            ui.output_plot("edad_hist"),
            ui.output_plot("posicion_bar"),
            ui.output_plot("nacionalidad_pie")
    ) ) ) 

def server(input, output, session):
    def cargar_links():
        workbook = client.open_by_key(SHEET_ID)
        worksheet = workbook.worksheet(LINK_SHEET_NAME)
        
        values = worksheet.get_all_values()
        if len(values) < 2:
            return {}

        df_links = pd.DataFrame(values[1:], columns=values[0])
        historial_links = {}
        
        for _, row in df_links.iterrows():
            jugador = row["Jugador"]
            enlace = row["Enlace"]
            
            if jugador not in historial_links:
                historial_links[jugador] = []
            historial_links[jugador].append(enlace)

        return historial_links

    historial_links_data = reactive.Value(cargar_links())

    @reactive.effect
    @reactive.event(input.guardar_enlace)
    def guardar_enlace():
        jugador = input.jugador_detalle()
        enlace = input.nuevo_enlace()

        if not jugador or not enlace:
            return  

        # Obtener historial actual de enlaces
        historial_actual = historial_links_data.get().copy()
        historial_actual.setdefault(jugador, []).append(enlace)

        # Guardar en Google Sheets
        workbook = client.open_by_key(SHEET_ID)
        worksheet = workbook.worksheet(LINK_SHEET_NAME)
        
        worksheet.append_row([jugador, enlace])

        # Actualizar estado reactivo
        historial_links_data.set(historial_actual)

        # Limpiar el campo de texto
        ui.update_text("nuevo_enlace", value="")

    @output
    @render.ui
    def historial_enlaces():
        jugador = input.jugador_detalle()

        if not jugador:
            return HTML("<p>Selecciona un jugador para ver sus enlaces guardados.</p>")

        enlaces = historial_links_data.get().get(jugador, [])

        if not enlaces:
            return HTML("<p>No hay enlaces guardados para este jugador.</p>")

        enlaces_html = "<ul>" + "".join(f'<li><a href="{link}" target="_blank">{link}</a></li>' for link in enlaces) + "</ul>"

        return HTML(f"<h4>Historial de Enlaces</h4>{enlaces_html}")
    
    @reactive.poll(cargar_historial,interval_secs=20)
    def actualizar_historial():
        historial_data.set(cargar_historial())

    @output
    @render.plot
    def edad_hist():
        df = df_jugadores
        fig, ax = plt.subplots(figsize=(6, 4))
        sns.histplot(df["Edad"], bins=5, kde=True, ax=ax, color="blue")
        ax.set_title("Distribución de Edades", fontsize=12, fontweight="bold")
        ax.set_xlabel("Edad")
        ax.set_ylabel("Cantidad de Jugadores")
        return fig

    @output
    @render.plot
    def posicion_bar():
        df = df_jugadores
        fig, ax = plt.subplots(figsize=(6, 4))
        sns.countplot(y=df["position_1"], order=df["position_1"].value_counts().index, ax=ax, palette="viridis")
        ax.set_title("Reparto de Posiciones", fontsize=12, fontweight="bold")
        ax.set_xlabel("Cantidad de Jugadores")
        ax.set_ylabel("Posición")
        return fig

    @output
    @render.plot
    def nacionalidad_pie():
        df = df_jugadores
        nacionalidades_expandidas = df["Nationality"].str.split(",").explode()
        nacionalidad_counts = nacionalidades_expandidas.value_counts()

        top_n = input.top_n()
        top_nacionalidades = nacionalidad_counts.nlargest(top_n)
        otros = nacionalidad_counts.iloc[top_n:].sum()

        data_pie = top_nacionalidades.to_dict()
        if otros > 0:
            data_pie["Otros"] = otros

        fig, ax = plt.subplots(figsize=(6, 6))
        colors = sns.color_palette("Set2", len(data_pie))
        explode = [0.1] + [0] * (len(data_pie) - 1)

        wedges, texts, autotexts = ax.pie(
            data_pie.values(),
            labels=data_pie.keys(),
            autopct="%1.1f%%",
            startangle=90,
            colors=colors,
            explode=explode,
            wedgeprops={"edgecolor": "black"}
        )

        for text in texts + autotexts:
            text.set_fontsize(10)
            text.set_fontweight("bold")

        ax.set_title("Distribución de Jugadores por Nacionalidad\n(Top {} + Otros)".format(top_n), fontsize=12, fontweight="bold")
        return fig
    
    @reactive.effect
    @reactive.event(input.jugador_detalle)
    def actualizar_opciones_informes():
        jugador = input.jugador_detalle()

        if not jugador:
            ui.update_select("informe_seleccionado", choices=["Nuevo Informe"])
            return

        informes = historial_data.get().get(jugador, [])
        opciones_informes = ["Nuevo Informe"] + [i["Fecha"] + " - " + i["Título"] for i in informes]

        ui.update_select("informe_seleccionado", choices=opciones_informes, selected="Nuevo Informe")

    @reactive.effect
    @reactive.event(input.restablecer_filtros)
    def restablecer_filtros():
        ui.update_text("filtro_nombre", value="")
        ui.update_text("filtro_comentarios", value="")
        ui.update_select("filtro_posicion1", selected="")
        ui.update_select("filtro_posicion2", selected="")
        ui.update_select("filtro_equipo", selected="")
        ui.update_select("filtro_nacionalidad", selected="")
        ui.update_select("filtro_agencia", selected="")
        ui.update_slider("filtro_edad_min", value=18)
        ui.update_slider("filtro_edad_max", value=25)
        ui.update_select("filtro_seguimiento", selected="")
        ui.update_select("filtro_categoria", selected="")

    @reactive.effect
    @reactive.poll(lambda: cargar_datos()[1], interval_secs=20) 
    def actualizar_datos_reactivos():
        nuevo_df, _ = cargar_datos()  # Cargar los datos más recientes
        df_reactivo.set(nuevo_df)  # Actualizar el DataFrame reactivo
    
    @reactive.effect
    @reactive.event(input.actualizar)
    def actualizar_datos():
        df = df_reactivo().copy()
        jugador = input.jugador()
        columna = input.columna()
        nuevo_valor = input.nuevo_valor()

        if jugador in df["full_name"].values and columna in df.columns:
            # Obtener el índice de la fila y la columna en Google Sheets
            col_index = df.columns.get_loc(columna) + 1  # Índice de columna en Google Sheets (1-based)
            row_index = df[df["full_name"] == jugador].index[0] + 2  # Fila en Google Sheets (1-based)

            try:
                worksheet.update_cell(row_index, col_index, nuevo_valor)  # Actualizar solo la celda necesaria
                df_reactivo()
                
                ui.notification_show(
                f"'{columna}' de {jugador} actualizado a '{nuevo_valor}'",
                type="message",
                duration=3
            )

            except Exception as e:
                # ❌ Notificación de error
                ui.notification_show(
                    "Error actualizando Google Sheets.",
                    type="error",
                    duration=3
                )
                print(f"Error actualizando Google Sheets: {e}") 
  
    @reactive.effect
    @reactive.event(input.ordenar)
    def actualizar_orden():
        df = df_reactivo().copy()
        columna = input.orden_columna()
        direccion = input.orden_direccion()

        if columna in df.columns:
            df = df.sort_values(by=columna, ascending=(direccion == "asc"))

        df_reactivo.set(df)

    @output
    @render.ui
    def tabla_actualizada():
        df = df_reactivo().copy()

        # Aplicar filtros
        if input.filtro_nombre():
            df = df[df["full_name"].str.contains(input.filtro_nombre(), case=False, na=False)]
        if input.filtro_comentarios():
            df = df[df["Comentarios"].str.contains(input.filtro_comentarios(), case=False, na=False)]
        if input.filtro_posicion1():
            df = df[df["position_1"] == input.filtro_posicion1()]
        if input.filtro_posicion2():
            df = df[df["position_2"] == input.filtro_posicion2()]
        if input.filtro_equipo():
            df = df[df["Team"] == input.filtro_equipo()]
        if input.filtro_nacionalidad():
            df = df[df["Nationality"].str.contains(input.filtro_nacionalidad(), na=False)]
        if "Agency" in df.columns and input.filtro_agencia():
            df = df[df["Agency"] == input.filtro_agencia()]
        if "Assessment" in df.columns and input.filtro_seguimiento():
            df = df[df["Assessment"] == input.filtro_seguimiento()]
        if "Categoria" in df.columns and input.filtro_categoria():
            df = df[df["Categoria"] == input.filtro_categoria()]
        if input.doble_nacionalidad():
            df = df[df["Nationality"].apply(es_doble_nacionalidad)]

        if "year_of_birth" in df.columns:
            edad_min = input.filtro_edad_min()
            edad_max = input.filtro_edad_max()
            df = df[(df["year_of_birth"] >= edad_min) & (df["year_of_birth"] <= edad_max)]

        # Función para colorear celdas y aplicar negrita al nombre
        def colorize(col, value):
            if col == "full_name":
                return f'<td><strong>{value}</strong></td>'  # Nombre en negrita
            elif col == "Contacto" and value in contacto_colors:
                return f'<td style="{contacto_colors[value]}">{value}</td>'
            elif value in assessment_colors:
                return f'<td style="{assessment_colors.get(value, "")}">{value}</td>'
            else:
                return f'<td>{value}</td>'

            # Generar tabla HTML con estilos mejorados
        table_html = '''
        <style>
            table.dataframe {
                border-collapse: collapse;
                width: 100%;
                font-size: 12px;  /* Reducir el tamaño del texto */
                margin-bottom: 20px;
                max-height: 800px;  /* Altura máxima para que no ocupe mucho espacio */
                overflow-y: auto;  /* Hacer que la tabla tenga scroll si es necesario */
                display: block;
            }
            table.dataframe th, table.dataframe td {
                border: 1px solid #ddd;
                padding: 4px;  /* Reducir padding para que la tabla sea más compacta */
                text-align: left;
            }
            table.dataframe th {
                background-color: #f2f2f2;
                color: #333;
            }
            table.dataframe tr:nth-child(even) {
                background-color: #f9f9f9;
            }
            table.dataframe tr:hover {
                background-color: #f1f1f1;
            }
        </style>
        <script>
            document.addEventListener("DOMContentLoaded", function() {
                document.querySelectorAll(".jugador-link").forEach(function(link) {
                    link.addEventListener("click", function(event) {
                        event.preventDefault();
                        let jugador = this.getAttribute("data-jugador");
                        Shiny.setInputValue("jugador_seleccionado", jugador, {priority: "event"});
                    });
                });
            });
        </script>
        <table class="dataframe">
        '''
        table_html += "<thead><tr>" + "".join(f"<th>{col}</th>" for col in df.columns) + "</tr></thead>"
        table_html += "<tbody>"
        for _, row in df.iterrows():
            table_html += "<tr>" + "".join(colorize(col, value) for col, value in row.items()) + "</tr>"
        table_html += "</tbody></table>"

        return HTML(table_html)

    @reactive.effect
    @reactive.event(input.agregar_jugador)
    def agregar_nuevo_jugador():
        df = df_reactivo().copy()
        
        nuevo_jugador = {
            "full_name": input.nuevo_nombre(),
            "position_1": input.nueva_posicion1(),
            "position_2": input.nueva_posicion2(),
            "Birthdate": str(input.nuevo_birthdate()),  # Convertir a string si es un objeto fecha
            "Foot": input.nuevo_pie(),
            "Team": input.nuevo_equipo(),
            "Assessment": input.nuevo_assessment(),
            "Nationality": input.nueva_nacionalidad(),
            "Agency": input.nueva_agencia(),
            "Club Contract": input.nuevo_club_contract(),
            "Contacto": input.nuevo_contacto(),
            "Categoria": input.nueva_categoria(),
            "Comentario": input.nuevo_comentario()
        }

        try:
            # Agregar a Google Sheets
            worksheet.append_row(list(nuevo_jugador.values()))

            values = worksheet.get_all_values()
            df_reactivo()
            
            # ✅ Notificación de éxito
            ui.notification_show(
                f"✅ {nuevo_jugador['full_name']} agregado correctamente",
                type="message",
                duration=3
            )
            
        except Exception as e:
            print(f"Error al agregar jugador: {e}")  # Capturar errores para debug

    @reactive.effect
    @reactive.event(input.eliminar_informe)
    def eliminar_informe():
        jugador = input.jugador_detalle()
        informe_titulo = input.informe_seleccionado()

        if not jugador or informe_titulo == "Nuevo Informe":
            return  # No hacer nada si no hay jugador o si no se ha seleccionado un informe válido

        historial_actual = historial_data.get().copy()
        informes = historial_actual.get(jugador, [])

        # Filtrar para eliminar el informe seleccionado
        informes = [i for i in informes if i["Fecha"] + " - " + i["Título"] != informe_titulo]

        # Si ya no hay informes para el jugador, eliminar la entrada
        if informes:
            historial_actual[jugador] = informes
        else:
            historial_actual.pop(jugador, None)

        guardar_historial(historial_actual)
        historial_data.set(historial_actual)

        # 🔹 Actualizar el selector para que vuelva a "Nuevo Informe"
        opciones_informes = ["Nuevo Informe"] + [i["Fecha"] + " - " + i["Título"] for i in informes]
        ui.update_select("informe_seleccionado", choices=opciones_informes, selected="Nuevo Informe")

        # 🔹 Limpiar los campos de texto
        ui.update_text("titulo_informe", value="")
        ui.update_text("texto_informe", value="")
        
            # ✅ Notificación de éxito
        ui.notification_show(
            f"📄 '{informe_titulo}' eliminado correctamente.",
            type="warning",  # Amarillo para advertencia/salida de datos
            duration=3
        )

        # Diccionario reactivo para almacenar informes por jugador
    historial_data = reactive.Value(cargar_historial())
    
    @output
    @render.ui
    def lapreferente_jugador():
        jugador = input.jugador_detalle()

        if not jugador:
            return HTML("Selecciona un jugador para ver su información en La Preferente.")

        df = df_reactivo().copy()
        
        # Obtener el equipo del jugador si la columna existe
        equipo = df[df["full_name"] == jugador]["Team"].iloc[0] if "Team" in df.columns and not df[df["full_name"] == jugador].empty else ""

        # Buscar enlace en La Preferente
        link = buscar_equipo_jugador(jugador, equipo)

        # Construir HTML
        if link:
            contenido = f'<p><a href="{link}" target="_blank">{link}</a></p>'
        else:
            contenido = "<p>No se encontró ningún resultado en La Preferente para este jugador.</p>"

        # Crear tarjeta con el enlace
        tarjeta_preferente = generar_tarjeta("La Preferente", {"Primer resultado": contenido})

        return HTML(tarjeta_preferente)

    @output
    @render.ui
    def seleccionar_informe():
        jugador = input.jugador_detalle()
        
        if not jugador:
            return HTML("<p>Selecciona un jugador para elegir un informe.</p>")
        
        informes = historial_data.get().get(jugador, [])
        
        if not informes:
            return HTML("<p>No hay informes disponibles para editar.</p>")
        
        opciones_informes = {i["Fecha"] + " - " + i["Título"]: i for i in informes}
        
        return ui.input_select("informe_seleccionado", "Seleccionar Informe", choices=list(opciones_informes.keys()), multiple=False)

    @reactive.effect
    @reactive.event(input.guardar_informe)
    def guardar_informe():
        jugador = input.jugador_detalle()
        titulo = input.titulo_informe()
        texto = input.texto_informe()
        informe_titulo = input.informe_seleccionado()

        if not jugador or not titulo or not texto:
            return  

        historial_actual = historial_data.get().copy()
        informes = historial_actual.get(jugador, [])

        # Verifica si estamos editando un informe existente
        for informe in informes:
            if informe["Fecha"] + " - " + informe["Título"] == informe_titulo:
                informe["Título"] = titulo
                informe["Texto"] = texto
                mensaje = f"📄 Informe '{titulo}' actualizado correctamente."
                break
        else:
            # Si no hay edición, agrega un nuevo informe
            fecha = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            informes.append({"Fecha": fecha, "Título": titulo, "Texto": texto})
            mensaje = f"✅ Informe '{titulo}' guardado correctamente."

        historial_actual[jugador] = informes
        guardar_historial(historial_actual)
        historial_data.set(historial_actual)

        ui.update_text("titulo_informe", value="")
        ui.update_text("texto_informe", value="")
            # 🔹 Actualizar el select para que vuelva a "Nuevo Informe"
        opciones_informes = ["Nuevo Informe"] + [i["Fecha"] + " - " + i["Título"] for i in informes]
        ui.update_select("informe_seleccionado", choices=opciones_informes, selected="Nuevo Informe")
        
        ui.notification_show(
        mensaje,
        type="message",  # Azul para información
        duration=3
        )

    @output
    @render.download
    def descargar_historial():
        jugador_seleccionado = input.jugador_detalle()

        if not jugador_seleccionado:
            return None

        informes = historial_data.get().get(jugador_seleccionado, [])
        
        if not informes:
            ui.notification_show(
            "⚠️ No hay informes disponibles para descargar.",
            type="warning",
            duration=3
            )
            return None

        doc = Document()
        doc.add_heading(f'Historial de Informes - {jugador_seleccionado}', level=1)

        for informe in informes:
            doc.add_paragraph(f"📅 Fecha: {informe['Fecha']}", style="Normal")
            doc.add_paragraph(f"📝 Título: {informe['Título']}", style="Normal")
            doc.add_paragraph(f"{informe['Texto']}", style="Normal")
            doc.add_paragraph("\n" + "-"*40 + "\n")

        file_path = f"historial_{jugador_seleccionado}.docx"
        doc.save(file_path)
        
        ui.notification_show(
        f"📥 Historial de informes de {jugador_seleccionado} generado.",
        type="message",
        duration=3
        )

        return file_path

    @output
    @render.ui
    def detalle_jugador():
        jugador = input.jugador_detalle()

        if not jugador:
            return HTML("<p>Selecciona un jugador para ver los detalles.</p>")

        df = df_reactivo().copy()
        jugador_data = df[df["full_name"] == jugador]

        if jugador_data.empty:
            return HTML("<p>No se encontraron datos para este jugador.</p>")

        jugador_data = jugador_data.iloc[0]

        # Información Personal
        info_personal = {
            "Nombre": jugador,
            "Año de Nacimiento": str(jugador_data.get("year_of_birth", "Desconocido"))[:4],
            "Nacionalidad": jugador_data.get("Nationality", "No disponible")
        }

        # Datos Deportivos
        info_deportiva = {
            "Posición Principal": jugador_data.get("position_1", "No disponible"),
            "Posición Secundaria": jugador_data.get("position_2", "No disponible"),
            "Equipo": jugador_data.get("Team", "Sin equipo"),
            "Agencia": jugador_data.get("Agency", "Sin agencia"),
            "Año Fin de Contrato": str(jugador_data.get("Año Fin Contrato", "No disponible"))
        }

        # ✅ Historial de Informes (Corrigiendo duplicación)
        historial = historial_data.get().get(jugador, [])
        informes_texto = "<p>No hay informes disponibles.</p>"
        if historial:
            informes_texto = "<ul>"
            for informe in historial:
                informes_texto += f"<li><strong>{informe['Fecha']} - {informe['Título']}</strong><br>{informe['Texto']}</li>"
            informes_texto += "</ul>"

        # Generar las tarjetas
        tarjeta_personal = generar_tarjeta("Información Personal", info_personal)
        tarjeta_deportiva = generar_tarjeta("Datos Deportivos", info_deportiva)
        tarjeta_informes = generar_tarjeta("Historial de Informes", {"Informes": informes_texto})

        # Combinar las tarjetas en un solo HTML
        detalles_html = f'''
        <div style="display: flex; flex-direction: column; gap: 20px;">
            {tarjeta_personal}
            {tarjeta_deportiva}
            {tarjeta_informes}
        </div>
        '''

        return HTML(detalles_html)

app = App(app_ui, server)